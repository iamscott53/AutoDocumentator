"""Document generator - exports recordings to HTML, Markdown, and DOCX."""

import base64
import shutil
from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from config import TEMPLATES_DIR, OUTPUT_DIR
from src.models import ActionType, Recording


class DocumentGenerator:
    """Generates documentation in various formats from a Recording."""

    def __init__(self):
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR)),
            autoescape=True,  # Escape user data to prevent XSS
        )

    def export_html(
        self,
        recording: Recording,
        output_path: Path | None = None,
        embed_images: bool = True,
    ) -> Path:
        """Export recording as a self-contained HTML document.

        Args:
            recording: The recording to export
            output_path: Where to save the HTML file
            embed_images: If True, embed images as base64 data URIs

        Returns:
            Path to the generated HTML file
        """
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = OUTPUT_DIR / f"SOP_{timestamp}.html"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Prepare step data with embedded images
        steps_data = []
        for step in recording.steps:
            step_dict = {
                "number": step.number,
                "action_type": step.action_type.value,
                "description": step.get_description(),
                "typed_text": step.typed_text,
                "hotkey_combo": step.hotkey_combo,
                "window_title": step.window_title,
                "image_src": "",
            }

            # Get the best available image
            img_path = step.thumbnail_path or step.annotated_screenshot_path
            if img_path and img_path.exists():
                try:
                    if embed_images:
                        step_dict["image_src"] = self._image_to_data_uri(img_path)
                    else:
                        img_dir = output_path.parent / "images"
                        img_dir.mkdir(exist_ok=True)
                        dest = img_dir / img_path.name
                        shutil.copy2(img_path, dest)
                        step_dict["image_src"] = f"images/{img_path.name}"
                except OSError:
                    pass  # Skip image if file became inaccessible

            # Full screenshot link (only as separate files, never embedded)
            full_img = step.annotated_screenshot_path or step.screenshot_path
            if full_img and full_img.exists() and not embed_images:
                try:
                    img_dir = output_path.parent / "images"
                    img_dir.mkdir(exist_ok=True)
                    dest = img_dir / f"full_{full_img.name}"
                    shutil.copy2(full_img, dest)
                    step_dict["full_image_src"] = f"images/full_{full_img.name}"
                except OSError:
                    step_dict["full_image_src"] = ""
            else:
                step_dict["full_image_src"] = ""

            steps_data.append(step_dict)

        template = self.jinja_env.get_template("sop_template.html")
        html_content = template.render(
            title=recording.title,
            description=recording.description,
            steps=steps_data,
            total_steps=len(steps_data),
            generated_date=datetime.now().strftime("%B %d, %Y at %I:%M %p"),
            duration=self._format_duration(recording.duration),
        )

        output_path.write_text(html_content, encoding="utf-8")
        return output_path

    def export_markdown(
        self,
        recording: Recording,
        output_path: Path | None = None,
    ) -> Path:
        """Export recording as a Markdown document with linked images.

        Args:
            recording: The recording to export
            output_path: Where to save the Markdown file

        Returns:
            Path to the generated Markdown file
        """
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = OUTPUT_DIR / f"SOP_{timestamp}.md"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create images directory next to the markdown file
        img_dir = output_path.parent / f"{output_path.stem}_images"
        img_dir.mkdir(exist_ok=True)

        lines = []
        lines.append(f"# {recording.title}\n")

        if recording.description:
            lines.append(f"{recording.description}\n")

        lines.append(f"**Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        lines.append(f"**Total Steps:** {len(recording.steps)}")
        lines.append(f"**Duration:** {self._format_duration(recording.duration)}\n")
        lines.append("---\n")

        for step in recording.steps:
            description = step.get_description()
            lines.append(f"## Step {step.number}: {description}\n")

            # Copy and reference the image
            img_path = step.thumbnail_path or step.annotated_screenshot_path
            if img_path and img_path.exists():
                dest = img_dir / img_path.name
                shutil.copy2(img_path, dest)
                rel_path = f"{img_dir.name}/{img_path.name}"
                lines.append(f"![Step {step.number}]({rel_path})\n")

            # Add details
            if step.typed_text:
                lines.append(f"**Text entered:** `{step.typed_text}`\n")
            if step.hotkey_combo:
                lines.append(f"**Keys pressed:** `{step.hotkey_combo}`\n")
            if step.window_title:
                lines.append(f"**Application:** {step.window_title}\n")

            lines.append("")

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return output_path

    def export_docx(
        self,
        recording: Recording,
        output_path: Path | None = None,
    ) -> Path:
        """Export recording as a Word document.

        Args:
            recording: The recording to export
            output_path: Where to save the DOCX file

        Returns:
            Path to the generated DOCX file
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = OUTPUT_DIR / f"SOP_{timestamp}.docx"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title
        title = doc.add_heading(recording.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if recording.description:
            doc.add_paragraph(recording.description)

        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
        meta.add_run(f"Total Steps: {len(recording.steps)}\n")
        meta.add_run(f"Duration: {self._format_duration(recording.duration)}")

        doc.add_paragraph()  # Spacer

        # Steps
        for step in recording.steps:
            description = step.get_description()

            # Step heading
            heading = doc.add_heading(f"Step {step.number}", level=2)

            # Description
            desc_para = doc.add_paragraph()
            run = desc_para.add_run(description)
            run.bold = True
            run.font.size = Pt(12)

            # Screenshot
            img_path = step.thumbnail_path or step.annotated_screenshot_path
            if img_path and img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph("[Screenshot could not be embedded]")

            # Details
            if step.typed_text:
                p = doc.add_paragraph()
                p.add_run("Text entered: ").bold = True
                p.add_run(step.typed_text)

            if step.hotkey_combo:
                p = doc.add_paragraph()
                p.add_run("Keys pressed: ").bold = True
                p.add_run(step.hotkey_combo)

            doc.add_paragraph()  # Spacer

        doc.save(str(output_path))
        return output_path

    @staticmethod
    def _image_to_data_uri(image_path: Path) -> str:
        """Convert an image file to a base64 data URI."""
        with open(image_path, "rb") as f:
            data = base64.b64encode(f.read()).decode("utf-8")
        return f"data:image/png;base64,{data}"

    @staticmethod
    def _format_duration(seconds: float) -> str:
        """Format a duration in seconds to a readable string."""
        if seconds < 60:
            return f"{int(seconds)} seconds"
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        if minutes < 60:
            return f"{minutes}m {secs}s"
        hours = minutes // 60
        mins = minutes % 60
        return f"{hours}h {mins}m"
