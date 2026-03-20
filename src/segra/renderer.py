"""Renders a validated SOPDocument to HTML, Markdown, or DOCX."""

from datetime import datetime
from pathlib import Path

from src.segra.schemas import SOPDocument


def render_html(sop: SOPDocument) -> str:
    """Render an SOPDocument to a self-contained HTML string."""
    steps_html = ""
    for s in sop.procedure_steps:
        expected = f'<div class="expected">Expected: {_esc(s.expected_result)}</div>' if s.expected_result else ""
        steps_html += (
            f'<tr><td class="step-num">{s.step}</td>'
            f'<td>{_esc(s.action)}{expected}</td></tr>\n'
        )

    prereqs = "".join(f"<li>{_esc(p)}</li>" for p in sop.prerequisites)
    validation = "".join(f"<li>{_esc(v)}</li>" for v in sop.validation)
    rollback = "".join(f"<li>{_esc(r)}</li>" for r in sop.rollback)
    security = "".join(f"<li>{_esc(n)}</li>" for n in sop.security_notes)
    refs = "".join(f"<li>{_esc(r)}</li>" for r in sop.references)

    troubleshooting = ""
    for t in sop.troubleshooting:
        troubleshooting += (
            f"<tr><td>{_esc(t.symptom)}</td>"
            f"<td>{_esc(t.cause)}</td>"
            f"<td>{_esc(t.fix)}</td></tr>\n"
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{_esc(sop.title)}</title>
<style>
body {{ font-family: -apple-system, 'Segoe UI', Roboto, sans-serif; max-width: 900px; margin: 40px auto; padding: 0 20px; color: #1a1a2e; }}
h1 {{ border-bottom: 3px solid #0f3460; padding-bottom: 10px; }}
h2 {{ color: #0f3460; margin-top: 30px; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
th {{ background: #f0f0f0; font-weight: 600; }}
.step-num {{ width: 50px; text-align: center; font-weight: 700; color: #e94560; }}
.expected {{ color: #666; font-size: 0.9em; margin-top: 5px; font-style: italic; }}
.meta {{ color: #666; font-size: 0.9em; margin-bottom: 20px; }}
.section-empty {{ color: #999; font-style: italic; }}
ul {{ margin: 5px 0; }}
@media print {{ body {{ max-width: 100%; }} }}
</style>
</head>
<body>
<h1>{_esc(sop.title)}</h1>
<p class="meta">Generated {datetime.now().strftime("%B %d, %Y at %I:%M %p")} by AutoDocumentator</p>
<h2>Purpose</h2><p>{_esc(sop.purpose)}</p>
{"<h2>Scope</h2><p>" + _esc(sop.scope) + "</p>" if sop.scope else ""}
{"<h2>Prerequisites</h2><ul>" + prereqs + "</ul>" if sop.prerequisites else ""}
<h2>Procedure</h2>
<table><tr><th>#</th><th>Action</th></tr>
{steps_html}
</table>
{"<h2>Validation</h2><ul>" + validation + "</ul>" if sop.validation else ""}
{"<h2>Rollback</h2><ul>" + rollback + "</ul>" if sop.rollback else ""}
{"<h2>Troubleshooting</h2><table><tr><th>Symptom</th><th>Cause</th><th>Fix</th></tr>" + troubleshooting + "</table>" if sop.troubleshooting else ""}
{"<h2>Security Notes</h2><ul>" + security + "</ul>" if sop.security_notes else ""}
{"<h2>References</h2><ul>" + refs + "</ul>" if sop.references else ""}
</body></html>"""


def render_markdown(sop: SOPDocument) -> str:
    """Render an SOPDocument to Markdown."""
    lines = [f"# {sop.title}\n"]
    lines.append(f"**Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
    lines.append(f"## Purpose\n\n{sop.purpose}\n")

    if sop.scope:
        lines.append(f"## Scope\n\n{sop.scope}\n")

    if sop.prerequisites:
        lines.append("## Prerequisites\n")
        for p in sop.prerequisites:
            lines.append(f"- {p}")
        lines.append("")

    lines.append("## Procedure\n")
    lines.append("| # | Action | Expected Result |")
    lines.append("|---|--------|-----------------|")
    for s in sop.procedure_steps:
        lines.append(f"| {s.step} | {s.action} | {s.expected_result} |")
    lines.append("")

    if sop.validation:
        lines.append("## Validation\n")
        for v in sop.validation:
            lines.append(f"- {v}")
        lines.append("")

    if sop.rollback:
        lines.append("## Rollback\n")
        for r in sop.rollback:
            lines.append(f"- {r}")
        lines.append("")

    if sop.troubleshooting:
        lines.append("## Troubleshooting\n")
        lines.append("| Symptom | Cause | Fix |")
        lines.append("|---------|-------|-----|")
        for t in sop.troubleshooting:
            lines.append(f"| {t.symptom} | {t.cause} | {t.fix} |")
        lines.append("")

    if sop.security_notes:
        lines.append("## Security Notes\n")
        for n in sop.security_notes:
            lines.append(f"- {n}")
        lines.append("")

    if sop.references:
        lines.append("## References\n")
        for r in sop.references:
            lines.append(f"- {r}")
        lines.append("")

    return "\n".join(lines)


def render_docx(sop: SOPDocument, output_path: Path) -> Path:
    """Render an SOPDocument to a Word document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading(sop.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_heading("Purpose", level=2)
    doc.add_paragraph(sop.purpose)

    if sop.scope:
        doc.add_heading("Scope", level=2)
        doc.add_paragraph(sop.scope)

    if sop.prerequisites:
        doc.add_heading("Prerequisites", level=2)
        for p in sop.prerequisites:
            doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Procedure", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Action"
    hdr[2].text = "Expected Result"
    for s in sop.procedure_steps:
        row = table.add_row().cells
        row[0].text = str(s.step)
        row[1].text = s.action
        row[2].text = s.expected_result

    if sop.validation:
        doc.add_heading("Validation", level=2)
        for v in sop.validation:
            doc.add_paragraph(v, style="List Bullet")

    if sop.rollback:
        doc.add_heading("Rollback", level=2)
        for r in sop.rollback:
            doc.add_paragraph(r, style="List Bullet")

    if sop.troubleshooting:
        doc.add_heading("Troubleshooting", level=2)
        t_table = doc.add_table(rows=1, cols=3)
        t_table.style = "Table Grid"
        t_hdr = t_table.rows[0].cells
        t_hdr[0].text = "Symptom"
        t_hdr[1].text = "Cause"
        t_hdr[2].text = "Fix"
        for t in sop.troubleshooting:
            row = t_table.add_row().cells
            row[0].text = t.symptom
            row[1].text = t.cause
            row[2].text = t.fix

    if sop.security_notes:
        doc.add_heading("Security Notes", level=2)
        for n in sop.security_notes:
            doc.add_paragraph(n, style="List Bullet")

    if sop.references:
        doc.add_heading("References", level=2)
        for r in sop.references:
            doc.add_paragraph(r, style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _esc(text: str) -> str:
    """HTML-escape a string (all five dangerous characters)."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#x27;")
    )
